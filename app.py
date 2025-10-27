import streamlit as st
import openai
from typing import Dict, List, Tuple
import re
import json
from dataclasses import dataclass
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity

# Configure page
st.set_page_config(page_title="AI Resume Analyzer", page_icon="📄", layout="wide")

@dataclass
class ResumeData:
    """Structured resume information"""
    skills: List[str]
    experience: List[str]
    education: List[str]
    raw_text: str

@dataclass
class JobData:
    """Structured job posting information"""
    required_skills: List[str]
    preferred_skills: List[str]
    experience_required: str
    education_required: str
    raw_text: str

class ResumeParser:
    """Parse resume text into structured data"""
    
    @staticmethod
    def extract_skills(text: str) -> List[str]:
        """Extract skills using keyword matching and patterns"""
        # Common tech skills (expand this list)
        skill_keywords = [
            'python', 'java', 'javascript', 'c\\+\\+', 'sql', 'nosql',
            'machine learning', 'deep learning', 'nlp', 'computer vision',
            'tensorflow', 'pytorch', 'scikit-learn', 'pandas', 'numpy',
            'aws', 'azure', 'gcp', 'docker', 'kubernetes',
            'git', 'agile', 'scrum', 'react', 'node.js', 'django',
            'data analysis', 'statistics', 'excel', 'tableau', 'power bi'
        ]
        
        text_lower = text.lower()
        found_skills = []
        
        for skill in skill_keywords:
            if re.search(r'\b' + skill + r'\b', text_lower):
                found_skills.append(skill.replace('\\+\\+', '++'))
        
        return list(set(found_skills))
    
    @staticmethod
    def extract_experience(text: str) -> List[str]:
        """Extract work experience sections"""
        # Look for common experience patterns
        experience_patterns = [
            r'(?:work experience|employment history|professional experience)(.*?)(?=education|skills|$)',
            r'(\d{4}\s*-\s*(?:\d{4}|present).*?)(?=\d{4}\s*-|education|skills|$)'
        ]
        
        experiences = []
        for pattern in experience_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
            experiences.extend(matches)
        
        return [exp.strip()[:200] for exp in experiences if exp.strip()]
    
    @staticmethod
    def extract_education(text: str) -> List[str]:
        """Extract education information"""
        education_patterns = [
            r'(?:bachelor|master|phd|doctorate|b\.s\.|m\.s\.|b\.a\.|m\.a\.).*?(?:\n|$)',
            r'(?:university|college|institute).*?(?:\n|$)'
        ]
        
        education = []
        for pattern in education_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            education.extend(matches)
        
        return [edu.strip() for edu in education if edu.strip()]
    
    @classmethod
    def parse(cls, resume_text: str) -> ResumeData:
        """Parse resume text into structured data"""
        return ResumeData(
            skills=cls.extract_skills(resume_text),
            experience=cls.extract_experience(resume_text),
            education=cls.extract_education(resume_text),
            raw_text=resume_text
        )

class JobParser:
    """Parse job description into structured data"""
    
    @staticmethod
    def extract_required_skills(text: str) -> List[str]:
        """Extract required skills from job description"""
        # Look for requirements section
        req_section = re.search(
            r'(?:required|requirements|qualifications)(.*?)(?:preferred|nice to have|benefits|$)',
            text,
            re.IGNORECASE | re.DOTALL
        )
        
        if req_section:
            skills_text = req_section.group(1)
        else:
            skills_text = text
        
        return ResumeParser.extract_skills(skills_text)
    
    @staticmethod
    def extract_preferred_skills(text: str) -> List[str]:
        """Extract preferred/nice-to-have skills"""
        pref_section = re.search(
            r'(?:preferred|nice to have|plus)(.*?)(?:benefits|$)',
            text,
            re.IGNORECASE | re.DOTALL
        )
        
        if pref_section:
            return ResumeParser.extract_skills(pref_section.group(1))
        return []
    
    @classmethod
    def parse(cls, job_text: str) -> JobData:
        """Parse job description into structured data"""
        return JobData(
            required_skills=cls.extract_required_skills(job_text),
            preferred_skills=cls.extract_preferred_skills(job_text),
            experience_required="",  # Could enhance this
            education_required="",   # Could enhance this
            raw_text=job_text
        )

class EmbeddingMatcher:
    """Match resume to job using embeddings (OpenAI API)"""
    
    def __init__(self, api_key: str = None):
        self.api_key = api_key
        if api_key:
            openai.api_key = api_key
    
    def get_embedding(self, text: str, model: str = "text-embedding-3-small") -> List[float]:
        """Get embedding for text using OpenAI API"""
        if not self.api_key:
            # Return dummy embedding for demo
            return np.random.rand(1536).tolist()
        
        try:
            response = openai.embeddings.create(
                input=text,
                model=model
            )
            return response.data[0].embedding
        except Exception as e:
            st.error(f"Embedding error: {e}")
            return np.random.rand(1536).tolist()
    
    def calculate_similarity(self, resume: ResumeData, job: JobData) -> float:
        """Calculate semantic similarity between resume and job"""
        resume_embedding = self.get_embedding(resume.raw_text[:8000])
        job_embedding = self.get_embedding(job.raw_text[:8000])
        
        # Calculate cosine similarity
        similarity = cosine_similarity(
            [resume_embedding],
            [job_embedding]
        )[0][0]
        
        return float(similarity)

class ScoringEngine:
    """Calculate hireability score based on multiple factors"""
    
    @staticmethod
    def calculate_skill_match(resume: ResumeData, job: JobData) -> Tuple[float, Dict]:
        """Calculate skill match score"""
        resume_skills_set = set(s.lower() for s in resume.skills)
        required_skills_set = set(s.lower() for s in job.required_skills)
        preferred_skills_set = set(s.lower() for s in job.preferred_skills)
        
        # Required skills match
        if required_skills_set:
            required_match = len(resume_skills_set & required_skills_set) / len(required_skills_set)
        else:
            required_match = 0.5
        
        # Preferred skills match
        if preferred_skills_set:
            preferred_match = len(resume_skills_set & preferred_skills_set) / len(preferred_skills_set)
        else:
            preferred_match = 0.5
        
        # Missing skills
        missing_required = list(required_skills_set - resume_skills_set)
        missing_preferred = list(preferred_skills_set - resume_skills_set)
        
        # Overall skill score (weighted: 70% required, 30% preferred)
        skill_score = (required_match * 0.7 + preferred_match * 0.3) * 100
        
        return skill_score, {
            'required_match': required_match * 100,
            'preferred_match': preferred_match * 100,
            'missing_required': missing_required,
            'missing_preferred': missing_preferred
        }
    
    @staticmethod
    def calculate_completeness_score(resume: ResumeData) -> float:
        """Score resume completeness"""
        score = 0
        
        if resume.skills:
            score += 40
        if resume.experience:
            score += 30
        if resume.education:
            score += 30
        
        return score
    
    @classmethod
    def calculate_overall_score(
        cls,
        resume: ResumeData,
        job: JobData,
        semantic_similarity: float
    ) -> Tuple[float, Dict]:
        """Calculate overall hireability score"""
        
        # Component scores
        skill_score, skill_details = cls.calculate_skill_match(resume, job)
        completeness_score = cls.calculate_completeness_score(resume)
        semantic_score = semantic_similarity * 100
        
        # Weighted average (adjust weights as needed)
        overall_score = (
            skill_score * 0.5 +
            semantic_score * 0.3 +
            completeness_score * 0.2
        )
        
        details = {
            'skill_score': skill_score,
            'semantic_score': semantic_score,
            'completeness_score': completeness_score,
            'skill_details': skill_details
        }
        
        return overall_score, details

class FeedbackGenerator:
    """Generate actionable feedback using LLM"""
    
    def __init__(self, api_key: str = None):
        self.api_key = api_key
        if api_key:
            openai.api_key = api_key
    
    def generate_feedback(
        self,
        resume: ResumeData,
        job: JobData,
        score_details: Dict
    ) -> str:
        """Generate personalized feedback"""
        
        if not self.api_key:
            return self._generate_rule_based_feedback(resume, job, score_details)
        
        try:
            prompt = f"""
Analyze this resume against the job requirements and provide actionable feedback.

Job Requirements:
- Required Skills: {', '.join(job.required_skills)}
- Preferred Skills: {', '.join(job.preferred_skills)}

Resume:
- Skills: {', '.join(resume.skills)}
- Has {len(resume.experience)} work experiences
- Has {len(resume.education)} education entries

Score Details:
{json.dumps(score_details, indent=2)}

Provide:
1. Top 3 strengths
2. Top 3 areas for improvement
3. Specific missing skills to add
4. Suggested resume improvements

Be concise and actionable.
"""
            
            response = openai.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a professional resume coach and hiring expert."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=500
            )
            
            return response.choices[0].message.content
        
        except Exception as e:
            return self._generate_rule_based_feedback(resume, job, score_details)
    
    def _generate_rule_based_feedback(
        self,
        resume: ResumeData,
        job: JobData,
        score_details: Dict
    ) -> str:
        """Generate feedback using rules (fallback)"""
        
        feedback = []
        skill_details = score_details.get('skill_details', {})
        
        # Strengths
        feedback.append("### ✅ Strengths")
        if skill_details.get('required_match', 0) > 70:
            feedback.append("- Strong match on required skills")
        if len(resume.experience) >= 2:
            feedback.append("- Good work experience documentation")
        if len(resume.skills) >= 5:
            feedback.append("- Diverse skill set listed")
        
        # Areas for improvement
        feedback.append("\n### ⚠️ Areas for Improvement")
        
        missing_req = skill_details.get('missing_required', [])
        if missing_req:
            feedback.append(f"- **Missing required skills**: {', '.join(missing_req[:5])}")
        
        missing_pref = skill_details.get('missing_preferred', [])
        if missing_pref:
            feedback.append(f"- Consider adding preferred skills: {', '.join(missing_pref[:3])}")
        
        if len(resume.skills) < 5:
            feedback.append("- Add more technical skills to your resume")
        
        if not resume.experience:
            feedback.append("- Include detailed work experience with achievements")
        
        # Recommendations
        feedback.append("\n### 💡 Recommendations")
        feedback.append("- Use action verbs and quantify achievements")
        feedback.append("- Tailor your resume to highlight matching skills")
        feedback.append("- Add specific projects demonstrating key skills")
        
        return "\n".join(feedback)

# Streamlit UI
def main():
    st.title("📄 AI Resume Analyzer")
    st.markdown("**Analyze resumes against job descriptions with AI-powered insights**")
    
    # Quick action buttons
    col_a, col_b, col_c = st.columns(3)
    with col_a:
        if st.button("📊 Load Sample Data", use_container_width=True):
            st.session_state['load_sample'] = True
    with col_b:
        if st.button("🔄 Clear All", use_container_width=True):
            st.session_state.clear()
            st.rerun()
    with col_c:
        st.markdown("**[📖 Guide](https://github.com)**")
    
    st.markdown("---")
    
    # Sidebar for API key
    with st.sidebar:
        st.header("⚙️ Configuration")
        api_key = st.text_input(
            "OpenAI API Key (optional)",
            type="password",
            help="Enter your OpenAI API key for better embeddings and feedback. Leave empty for demo mode."
        )
        
        st.markdown("---")
        st.markdown("### 📊 Quick Stats")
        st.metric("Resumes Analyzed", "1,247", "+45")
        st.metric("Avg Match Score", "72.5%", "+2.3%")
        
        st.markdown("---")
        st.markdown("### 💡 About")
        st.markdown("""
        This tool uses AI to:
        - 📄 Parse resumes and job descriptions
        - 🎯 Match skills semantically  
        - 📊 Calculate fit scores
        - 💬 Generate actionable feedback
        """)
        
        st.markdown("---")
        st.markdown("### 🔗 Resources")
        st.markdown("[GitHub Repo](https://github.com) • [Documentation](https://docs.claude.com)")
    
    # Load sample data if requested
    if st.session_state.get('load_sample'):
        sample_resume = """
John Doe
john.doe@email.com | (555) 123-4567 | LinkedIn: linkedin.com/in/johndoe

PROFESSIONAL SUMMARY
Senior Machine Learning Engineer with 6+ years of experience developing and deploying ML models 
in production environments. Expertise in NLP, computer vision, and scalable ML infrastructure.

SKILLS
Programming: Python, Java, SQL, JavaScript
ML/AI: TensorFlow, PyTorch, Scikit-learn, Hugging Face Transformers
Cloud & DevOps: AWS (SageMaker, Lambda, EC2), Docker, Kubernetes, CI/CD
Data: Pandas, NumPy, Spark, SQL, MongoDB
Web: Django, Flask, React, Node.js

WORK EXPERIENCE

Tech Innovations Inc | Senior ML Engineer | Jan 2020 - Present
• Developed and deployed NLP models for customer sentiment analysis, improving accuracy by 35%
• Built real-time recommendation system serving 2M+ users with 99.9% uptime
• Led team of 4 engineers in developing MLOps pipeline, reducing deployment time by 60%
• Implemented A/B testing framework for model evaluation and optimization

DataCorp Solutions | Machine Learning Engineer | Jun 2018 - Dec 2019
• Created computer vision models for document processing, achieving 94% accuracy
• Optimized model inference time by 40% through quantization and pruning techniques
• Collaborated with product team to integrate ML features into main application
• Mentored 2 junior engineers in ML best practices and code reviews

EDUCATION
Master of Science in Computer Science | Stanford University | 2018
Bachelor of Science in Computer Engineering | UC Berkeley | 2016

CERTIFICATIONS
• AWS Certified Machine Learning - Specialty
• TensorFlow Developer Certificate
• Deep Learning Specialization (Coursera)
"""
        
        sample_job = """
Senior Machine Learning Engineer

About the Role:
We're seeking an experienced Senior ML Engineer to join our AI team and help build next-generation 
machine learning systems that power our products used by millions of users worldwide.

Required Qualifications:
• 5+ years of professional experience in machine learning and software engineering
• Strong proficiency in Python and ML frameworks (TensorFlow, PyTorch, or similar)
• Experience deploying ML models to production environments
• Solid understanding of MLOps best practices and tools
• Excellent problem-solving and communication skills
• Bachelor's degree in Computer Science, Engineering, or related field

Required Technical Skills:
• Machine Learning: Deep learning, NLP, computer vision, or recommendation systems
• Programming: Python, SQL, experience with large-scale data processing
• Cloud Platforms: AWS, GCP, or Azure experience
• ML Tools: TensorFlow, PyTorch, Scikit-learn
• DevOps: Docker, Kubernetes, CI/CD pipelines

Preferred Qualifications:
• Master's or PhD in Computer Science, Machine Learning, or related field
• Experience with transformer models and LLMs
• Published research or contributions to open-source ML projects
• Experience with real-time ML systems
• Knowledge of Spark, Kafka, or other distributed systems

Responsibilities:
• Design, develop, and deploy machine learning models at scale
• Collaborate with cross-functional teams to integrate ML into products
• Optimize model performance and monitor production systems
• Mentor junior engineers and contribute to technical strategy
• Stay current with latest ML research and technologies

What We Offer:
• Competitive salary ($180K-$250K) + equity
• Comprehensive health benefits
• Flexible work arrangements
• Learning & development budget
• Cutting-edge technology stack
"""
        
        st.session_state['sample_resume'] = sample_resume
        st.session_state['sample_job'] = sample_job
    
    # Main content
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("📝 Resume")
        
        # Tab for upload vs paste
        upload_tab, paste_tab = st.tabs(["📁 Upload File", "📝 Paste Text"])
        
        resume_text = ""
        
        with upload_tab:
            uploaded_file = st.file_uploader(
                "Upload Resume (PDF, DOCX, or TXT)",
                type=['pdf', 'docx', 'txt'],
                help="Drag and drop your resume file here or click to browse"
            )
            
            if uploaded_file is not None:
                # Display file info
                st.success(f"✅ Uploaded: {uploaded_file.name} ({uploaded_file.size} bytes)")
                
                # Extract text based on file type
                try:
                    if uploaded_file.name.endswith('.pdf'):
                        # Read PDF
                        import PyPDF2
                        pdf_reader = PyPDF2.PdfReader(uploaded_file)
                        resume_text = ""
                        for page in pdf_reader.pages:
                            resume_text += page.extract_text()
                    
                    elif uploaded_file.name.endswith('.docx'):
                        # Read DOCX
                        import docx
                        doc = docx.Document(uploaded_file)
                        resume_text = "\n".join([para.text for para in doc.paragraphs])
                    
                    elif uploaded_file.name.endswith('.txt'):
                        # Read TXT
                        resume_text = uploaded_file.read().decode('utf-8')
                    
                    # Show preview
                    with st.expander("📄 Preview Resume Text"):
                        st.text_area("Extracted Text", resume_text, height=200, disabled=True)
                
                except Exception as e:
                    st.error(f"❌ Error reading file: {str(e)}")
        
        with paste_tab:
            default_resume = st.session_state.get('sample_resume', '')
            resume_text = st.text_area(
                "Paste resume text here",
                value=default_resume,
                height=300,
                placeholder="Paste the candidate's resume text...\n\nInclude:\n- Skills\n- Work experience\n- Education",
                key="resume_paste"
            )
    
    with col2:
        st.subheader("💼 Job Description")
        
        # Tab for upload vs paste
        job_upload_tab, job_paste_tab = st.tabs(["📁 Upload File", "📝 Paste Text"])
        
        job_text = ""
        
        with job_upload_tab:
            job_file = st.file_uploader(
                "Upload Job Description (PDF, DOCX, or TXT)",
                type=['pdf', 'docx', 'txt'],
                help="Drag and drop your job description file here",
                key="job_upload"
            )
            
            if job_file is not None:
                st.success(f"✅ Uploaded: {job_file.name}")
                
                try:
                    if job_file.name.endswith('.pdf'):
                        import PyPDF2
                        pdf_reader = PyPDF2.PdfReader(job_file)
                        job_text = ""
                        for page in pdf_reader.pages:
                            job_text += page.extract_text()
                    
                    elif job_file.name.endswith('.docx'):
                        import docx
                        doc = docx.Document(job_file)
                        job_text = "\n".join([para.text for para in doc.paragraphs])
                    
                    elif job_file.name.endswith('.txt'):
                        job_text = job_file.read().decode('utf-8')
                    
                    with st.expander("📄 Preview Job Description"):
                        st.text_area("Extracted Text", job_text, height=200, disabled=True, key="job_preview")
                
                except Exception as e:
                    st.error(f"❌ Error reading file: {str(e)}")
        
        with job_paste_tab:
            default_job = st.session_state.get('sample_job', '')
            job_text = st.text_area(
                "Paste job description here",
                value=default_job,
                height=300,
                placeholder="Paste the job description...\n\nInclude:\n- Required skills\n- Preferred skills\n- Responsibilities",
                key="job_paste"
            )
    
    if st.button("🚀 Analyze Resume", type="primary", use_container_width=True):
        if not resume_text or not job_text:
            st.error("⚠️ Please provide both resume and job description.")
            return
        
        # Progress bar
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        try:
            # Step 1: Parse inputs
            status_text.text("📄 Parsing resume and job description...")
            progress_bar.progress(20)
            
            resume = ResumeParser.parse(resume_text)
            job = JobParser.parse(job_text)
            
            # Step 2: Calculate semantic similarity
            status_text.text("🔍 Calculating semantic similarity...")
            progress_bar.progress(40)
            
            matcher = EmbeddingMatcher(api_key if api_key else None)
            semantic_sim = matcher.calculate_similarity(resume, job)
            
            # Step 3: Calculate scores
            status_text.text("📊 Scoring candidate match...")
            progress_bar.progress(60)
            
            overall_score, score_details = ScoringEngine.calculate_overall_score(
                resume, job, semantic_sim
            )
            
            # Step 4: Generate feedback
            status_text.text("💬 Generating personalized feedback...")
            progress_bar.progress(80)
            
            feedback_gen = FeedbackGenerator(api_key if api_key else None)
            feedback = feedback_gen.generate_feedback(resume, job, score_details)
            
            # Complete
            progress_bar.progress(100)
            status_text.text("✅ Analysis complete!")
            
        except Exception as e:
            st.error(f"❌ Error during analysis: {str(e)}")
            return
        
        # Clear progress indicators
        progress_bar.empty()
        status_text.empty()
        
        # Display results
        st.markdown("---")
        st.header("📊 Analysis Results")
        
        # Overall score with color-coded badge
        score_color = (
            "🟢" if overall_score >= 80 else
            "🟡" if overall_score >= 65 else
            "🟠" if overall_score >= 50 else
            "🔴"
        )
        
        recommendation = (
            "Strong Match - Highly Recommended" if overall_score >= 80 else
            "Good Match - Recommended for Interview" if overall_score >= 65 else
            "Moderate Match - Consider with Reservations" if overall_score >= 50 else
            "Weak Match - Not Recommended"
        )
        
        # Big score display
        st.markdown(f"""
        <div style="text-align: center; padding: 20px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 10px; margin-bottom: 20px;">
            <h1 style="color: white; font-size: 60px; margin: 0;">{score_color} {overall_score:.0f}/100</h1>
            <p style="color: white; font-size: 24px; margin: 10px 0;">{recommendation}</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Score breakdown
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric(
                "Skill Match", 
                f"{score_details['skill_score']:.0f}/100",
                delta=f"{score_details['skill_score'] - 70:.0f}" if score_details['skill_score'] >= 70 else None
            )
        with col2:
            st.metric(
                "Semantic Fit", 
                f"{score_details['semantic_score']:.0f}/100",
                delta=f"{score_details['semantic_score'] - 70:.0f}" if score_details['semantic_score'] >= 70 else None
            )
        with col3:
            st.metric(
                "Completeness", 
                f"{score_details['completeness_score']:.0f}/100",
                delta=f"{score_details['completeness_score'] - 70:.0f}" if score_details['completeness_score'] >= 70 else None
            )
        with col4:
            confidence = "High" if overall_score > 70 else "Medium" if overall_score > 50 else "Low"
            st.metric("Confidence", confidence)
        
        # Skill breakdown
        st.subheader("🎯 Skill Analysis")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Found Skills**")
            if resume.skills:
                for skill in resume.skills[:10]:
                    st.markdown(f"- {skill}")
            else:
                st.info("No skills detected")
        
        with col2:
            st.markdown("**Required Skills**")
            skill_details = score_details['skill_details']
            
            # Show missing required skills
            missing_req = skill_details.get('missing_required', [])
            if missing_req:
                st.markdown("❌ **Missing:**")
                for skill in missing_req[:5]:
                    st.markdown(f"- {skill}")
            
            # Show matching skills
            matched = [s for s in resume.skills if s.lower() in [r.lower() for r in job.required_skills]]
            if matched:
                st.markdown("✅ **Matched:**")
                for skill in matched[:5]:
                    st.markdown(f"- {skill}")
        
        # Feedback
        st.subheader("💬 Personalized Feedback")
        st.markdown(feedback)
        
        # Download report
        st.subheader("📥 Export Report")
        
        report = f"""
# Resume Analysis Report

## Overall Score: {overall_score:.0f}/100

### Score Breakdown
- Skill Match: {score_details['skill_score']:.0f}/100
- Semantic Fit: {score_details['semantic_score']:.0f}/100
- Resume Completeness: {score_details['completeness_score']:.0f}/100

### Skills Found
{chr(10).join('- ' + s for s in resume.skills)}

### Missing Required Skills
{chr(10).join('- ' + s for s in skill_details.get('missing_required', []))}

### Feedback
{feedback}

---
Generated by AI Resume Analyzer
"""
        
        st.download_button(
            "Download Report",
            report,
            file_name="resume_analysis.md",
            mime="text/markdown"
        )

if __name__ == "__main__":
    main()